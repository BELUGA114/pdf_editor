"""导出 — 同步更新/去红头 DOCX、格式保留合并"""

import tkinter as tk
from tkinter import filedialog
import difflib
from .config import _BaseMixin

class ExportMixin(_BaseMixin):
    """导出 — 同步更新/去红头 DOCX、格式保留合并"""

    def _merge_accepted_changes(self):
        """将已同意的更改写入 self.doc_obj（原地修改），返回并入的更改数量"""
        if not self.doc_obj:
            return 0
        accepted = [b for b in self.diff_blocks if b['accepted'] and b['tag'] != 'equal']
        if not accepted:
            return 0

        # 重建 accepted 目标 flat text（归一化空间）
        target_parts = []
        for b in self.diff_blocks:
            if b['tag'] == 'equal':
                target_parts.append(b['old'])
            elif b['tag'] == 'delete':
                if not b['accepted']:
                    target_parts.append(b['old'])
            elif b['tag'] == 'insert':
                if b['accepted']:
                    target_parts.append(b['new'])
            elif b['tag'] == 'replace':
                target_parts.append(b['new'] if b['accepted'] else b['old'])
        target_flat = ''.join(target_parts)

        # 在归一化空间做二次 diff，找到编辑操作
        full_text = getattr(self, '_saved_docx_source', self.docx_text)
        flat_to_full = self._docx_flat_positions
        original_flat = ''.join(full_text[i] for i in flat_to_full)

        matcher = difflib.SequenceMatcher(None, original_flat, target_flat)
        edits = []  # [(raw_start, raw_delete_len, insert_str), ...]
        raw_flat = self._docx_flat_to_raw
        raw_full = self._docx_text_raw
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                continue
            # 将 flat 空间的编辑位置映射到原始（未归一化）文本空间
            raw_start = raw_flat[i1] if i1 < len(raw_flat) else len(raw_full)
            raw_end = raw_flat[i2] if i2 < len(raw_flat) else len(raw_full)
            raw_dlen = raw_end - raw_start
            insert = target_flat[j1:j2] if tag in ('insert', 'replace') else ''
            edits.append((raw_start, raw_dlen, insert))

        # 逆序应用编辑到原始文本（保留全角符号、多空格和原始格式）
        edits.sort(key=lambda x: x[0], reverse=True)
        chars = list(raw_full)
        for start, dlen, insert in edits:
            del chars[start:start + dlen]
            for ch in reversed(insert):
                chars.insert(start, ch)
        modified = ''.join(chars)

        # 按段落边界回写到 DOCX
        new_paras = modified.split('\n')
        body_paras = [p for p in self.doc_obj.paragraphs
                      if not self._is_red_header(p) and p.text.strip()]
        if not body_paras:
            return len(accepted)
        for i, para_text in enumerate(new_paras):
            if i < len(body_paras):
                p = body_paras[i]
                if p.text == para_text:
                    continue
                runs = p.runs
                if not runs:
                    p.add_run(para_text)
                elif len(runs) == 1:
                    runs[0].text = para_text
                else:
                    old_lens = [len(r.text) for r in runs]
                    old_total = sum(old_lens)
                    new_len = len(para_text)
                    pos = 0
                    for j, r in enumerate(runs):
                        if j == len(runs) - 1:
                            r.text = para_text[pos:]
                        else:
                            if old_total > 0:
                                ratio = old_lens[j] / old_total
                                count = round(new_len * ratio)
                            else:
                                count = 0
                            remaining_runs = len(runs) - j - 1
                            count = max(count, 0)
                            count = min(count, new_len - pos - remaining_runs)
                            r.text = para_text[pos:pos + count] if count > 0 else ''
                            pos += count
            else:
                new_p = body_paras[-1].insert_paragraph_after(para_text)  # type: ignore[attr-defined]
                if new_p.runs:
                    new_p.runs[0].text = para_text
        for p in body_paras[len(new_paras):]:
            for r in p.runs:
                r.text = ''

        # 防止重复应用：清除 diff_blocks 及相关状态
        self.diff_blocks = []
        self._docx_flat_positions = []
        self._docx_flat_to_raw = []
        self._saved_docx_source = ''

        return len(accepted)



    def save_synced_docx(self):
        """导出：将已同意的更改并入 DOCX（未同意则保留原文）"""
        if not self.doc_obj:
            return
        save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                       filetypes=[("Word Documents", "*.docx")])
        if not save_path:
            return

        count = self._merge_accepted_changes()
        self.doc_obj.save(save_path)
        if count:
            self._alert("成功", f"已并入 {count} 处更改并导出文档。", "info")
        else:
            self._alert("成功", "未同意任何更改，已导出原版文档。", "info")



    def _is_red_border(self, paragraph):
        """检查段落是否有红色边框线（红头文件的红色横线）"""
        try:
            from docx.oxml.ns import qn
            pPr = paragraph._element.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    for tag in (qn('w:bottom'), qn('w:top')):
                        border = pBdr.find(tag)
                        if border is not None:
                            color = border.get(qn('w:color'))
                            if color and color != 'auto' and len(color) >= 6:
                                r = int(color[0:2], 16)
                                g = int(color[2:4], 16)
                                b = int(color[4:6], 16)
                                if r > 200 and g < 60 and b < 60:
                                    return True
        except Exception:
            pass
        return False



    def _remove_red_border(self, paragraph):
        """移除段落的红色边框线"""
        try:
            from docx.oxml.ns import qn
            pPr = paragraph._element.find(qn('w:pPr'))
            if pPr is not None:
                pBdr = pPr.find(qn('w:pBdr'))
                if pBdr is not None:
                    for tag in (qn('w:bottom'), qn('w:top')):
                        border = pBdr.find(tag)
                        if border is not None:
                            color = border.get(qn('w:color'))
                            if color and color != 'auto' and len(color) >= 6:
                                r = int(color[0:2], 16)
                                g = int(color[2:4], 16)
                                b = int(color[4:6], 16)
                                if r > 200 and g < 60 and b < 60:
                                    pBdr.remove(border)
        except Exception:
            pass



    def save_dered_docx(self):
        """导出：先并入已同意的更改，再去红头（清除红色字体段落和红色横线，不改变其余格式）"""
        if not self.doc_obj:
            return

        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if not save_path:
            return

        # 先并入已同意的更改
        self._merge_accepted_changes()

        # 扫描前6个段落，检测并清除红头元素
        paragraphs_to_check = self.doc_obj.paragraphs[:6]

        for p in paragraphs_to_check:
            # 检查红色文字
            is_red_head = False
            for run in p.runs:
                if run.font.color and run.font.color.rgb:
                    r, g, b = run.font.color.rgb
                    if r > 200 and g < 60 and b < 60:
                        is_red_head = True
                        break

            # 命中了红头特征（红色文字或包含"文件"字样），清除文本但保留格式
            if is_red_head or "文件" in p.text:
                # 逐个 run 清空文字（保留字体/大小等格式信息）
                for run in p.runs:
                    run.text = ''
                p.paragraph_format.space_before = 0
                p.paragraph_format.space_after = 0

            # 清除红色横线（段落边框）
            if self._is_red_border(p):
                self._remove_red_border(p)
                # 如果该段落只有边框没有文字，清空其内容
                if not p.text.strip():
                    for run in p.runs:
                        run.text = ''

        try:
            self.doc_obj.save(save_path)
            self._alert("成功", "已并入更改并去除红头，文档已导出。", "info")
        except Exception as e:
            self._alert("错误", f"保存失败: {str(e)}", "error")


