import tkinter as tk
from tkinter import filedialog
import os
import io
import ctypes
from ctypes import wintypes

import fitz  # PyMuPDF
import numpy as np
from rapidocr_onnxruntime import RapidOCR
from docx import Document
import difflib
from PIL import Image, ImageTk, ImageDraw

# ================= 配置区域 =================
ocr_engine = RapidOCR()
# ============================================

# Windows 拖放支持：注册窗口接受文件拖放消息
def _enable_dnd(hwnd: int, callback):
    """为 Windows 窗口启用文件拖放，拖入文件时调用 callback(paths)"""
    user32 = ctypes.windll.user32
    shell32 = ctypes.windll.shell32

    # 显式声明 API 签名，避免 64 位指针截断
    user32.GetWindowLongPtrW.argtypes = [wintypes.HWND, ctypes.c_int]
    user32.GetWindowLongPtrW.restype = wintypes.LPARAM
    user32.SetWindowLongPtrW.argtypes = [wintypes.HWND, ctypes.c_int, wintypes.LPARAM]
    user32.SetWindowLongPtrW.restype = wintypes.LPARAM
    user32.CallWindowProcW.argtypes = [wintypes.LPARAM, wintypes.HWND, wintypes.UINT,
                                       wintypes.WPARAM, wintypes.LPARAM]
    user32.CallWindowProcW.restype = wintypes.LPARAM

    shell32.DragAcceptFiles(hwnd, True)

    GWLP_WNDPROC = -4
    WM_DROPFILES = 0x0233

    WNDPROC = ctypes.WINFUNCTYPE(
        wintypes.LPARAM, wintypes.HWND, wintypes.UINT,
        wintypes.WPARAM, wintypes.LPARAM,
    )
    original = user32.GetWindowLongPtrW(hwnd, GWLP_WNDPROC)

    @WNDPROC
    def new_wndproc(hwnd, msg, wparam, lparam):
        if msg == WM_DROPFILES:
            hdrop = wparam
            count = shell32.DragQueryFileW(hdrop, 0xFFFFFFFF, None, 0)
            paths = []
            for i in range(count):
                buf = ctypes.create_unicode_buffer(260)
                shell32.DragQueryFileW(hdrop, i, buf, 260)
                paths.append(buf.value)
            shell32.DragFinish(hdrop)
            callback(paths)
            return 0
        return user32.CallWindowProcW(original, hwnd, msg, wparam, lparam)

    user32.SetWindowLongPtrW(hwnd, GWLP_WNDPROC,
                              ctypes.cast(new_wndproc, ctypes.c_void_p).value)
    return new_wndproc


class DocxPdfReviewer:
    # 颜色常量
    BG = "#f0f2f5"         # 页面背景
    CARD_BG = "#ffffff"    # 卡片背景
    PRIMARY = "#4a6eb5"    # 主色调
    SUCCESS = "#52c41a"    # 成功/导出
    WARN = "#faad14"       # 警告
    BORDER = "#d9d9d9"     # 边框
    TEXT = "#333333"       # 正文
    MUTED = "#999999"      # 弱化文字

    def __init__(self, root):
        self.root = root
        self.root.title("Docx-PDF 差异比对工具")
        self.root.geometry("1200x800")
        self.root.configure(bg=self.BG)
        # 全局字体
        self.root.option_add("*Font", ("Microsoft YaHei UI", 10))

        self.docx_path = None
        self.pdf_path = None
        self.doc_obj = None

        self.docx_text = ""
        self.pdf_text = ""
        self.diff_results = []
        self.cleaned_pdf_images = []
        self._saved_sel = None
        self.compare_symbols = tk.BooleanVar(value=True)
        self.show_info = tk.BooleanVar(value=True)
        self._docx_folder = None
        self._pdf_folder = None
        self._pairs = []           # 批量比对: [(docx_path, pdf_path), ...]
        self._pair_index = -1
        self._pair_data = {}       # {index: {docx_text, pdf_text, doc_obj, cleaned, ...}}
        self.diff_blocks = []      # [{tag, old, new, accepted, block_id}, ...]
        self._docx_paragraphs = [] # 原始段落文本（用于合并回写）
        self._hover_popup = None   # 悬浮窗引用
        self._hover_block_id = -1  # 当前悬浮的块ID

        self._build_ui()

        # 注册拖放（需在窗口实现后）
        self.root.update_idletasks()
        self._dnd_proc = _enable_dnd(self.root.winfo_id(), self._on_drop)

    def _on_drop(self, paths):
        for path in paths:
            if os.path.isdir(path):
                self._on_drop_folder(path)
            else:
                ext = os.path.splitext(path)[1].lower()
                if ext == ".docx":
                    self._load_docx_path(path)
                elif ext == ".pdf":
                    self._load_pdf_path(path)

    def _on_drop_folder(self, path):
        """拖入文件夹时自动判断类型（按文件数占比）"""
        docx_cnt = 0
        pdf_cnt = 0
        try:
            for f in os.listdir(path):
                ext = os.path.splitext(f)[1].lower()
                if ext in {'.docx', '.doc'}:
                    docx_cnt += 1
                elif ext == '.pdf':
                    pdf_cnt += 1
        except OSError:
            return

        if docx_cnt == 0 and pdf_cnt == 0:
            return

        if docx_cnt >= pdf_cnt:
            self._docx_folder = path
            self.lbl_docx_folder.config(text=os.path.basename(path), fg=self.PRIMARY)
        if pdf_cnt >= docx_cnt:
            self._pdf_folder = path
            self.lbl_pdf_folder.config(text=os.path.basename(path), fg=self.PRIMARY)

        self._try_folder_match()

    def _set_status(self, text: str, color: str = None):
        """更新操作栏状态标签"""
        self._status_lbl.config(text=text, fg=color or self.MUTED)
        self.root.update_idletasks()

    def _alert(self, title: str, message: str, level: str = "info"):
        """无提示音的消息弹窗（替代 messagebox）"""
        if level == "info" and not self.show_info.get():
            return
        colors = {"info": "#d1ecf1", "warn": "#fff3cd", "error": "#f8d7da"}
        win = tk.Toplevel(self.root)
        win.withdraw()
        win.title(title)
        win.resizable(False, False)
        win.transient(self.root)
        win.grab_set()
        bg = colors.get(level, "#d1ecf1")
        f = tk.Frame(win, padx=20, pady=15, bg=bg)
        f.pack(fill=tk.BOTH, expand=True)
        tk.Label(f, text=message, bg=bg, wraplength=400, justify=tk.LEFT,
                 font=("Microsoft YaHei UI", 10)).pack(pady=(0, 10))
        btn = tk.Button(f, text="确定", command=win.destroy, width=10,
                        bg=self.PRIMARY, fg="white", relief=tk.FLAT,
                        padx=20, pady=4, cursor="hand2")
        btn.pack()
        btn.focus_set()
        win.bind("<Return>", lambda _: win.destroy())
        win.bind("<Escape>", lambda _: win.destroy())
        win.update_idletasks()
        pw, ph = self.root.winfo_width(), self.root.winfo_height()
        px, py = self.root.winfo_x(), self.root.winfo_y()
        ww, wh = win.winfo_width(), win.winfo_height()
        x = px + (pw - ww) // 2
        y = py + (ph - wh) // 2
        win.geometry(f"+{x}+{y}")
        win.deiconify()
        win.wait_window()

    def _make_drop_zone(self, parent, title: str, subtitle: str, icon: str, on_click):
        """创建一个可点击/拖放的文件卡片区域（水平布局：图标 | 文字）"""
        card = tk.Frame(parent, bg=self.CARD_BG, highlightbackground=self.BORDER,
                        highlightthickness=1, cursor="hand2")
        inner = tk.Frame(card, bg=self.CARD_BG, padx=10, pady=6, cursor="hand2")
        inner.pack(fill=tk.BOTH, expand=True)

        def on_enter(_):
            card.config(highlightbackground=self.PRIMARY, highlightthickness=2)
        def on_leave(_):
            card.config(highlightbackground=self.BORDER, highlightthickness=1)

        widgets: list = [card, inner]

        # 左侧图标
        icon_lbl = tk.Label(inner, text=icon, font=("Segoe UI", 16), bg=self.CARD_BG,
                            cursor="hand2")
        icon_lbl.pack(side=tk.LEFT, padx=(0, 8))
        widgets.append(icon_lbl)

        # 右侧文字区域
        text_col = tk.Frame(inner, bg=self.CARD_BG, cursor="hand2")
        text_col.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        widgets.append(text_col)

        lbl_title = tk.Label(text_col, text=title, font=("Microsoft YaHei UI", 9, "bold"),
                             bg=self.CARD_BG, fg=self.TEXT, cursor="hand2",
                             anchor=tk.W)
        lbl_title.pack(fill=tk.X)
        widgets.append(lbl_title)

        sub_lbl = tk.Label(text_col, text=subtitle, font=("Microsoft YaHei UI", 7),
                           bg=self.CARD_BG, fg=self.MUTED, cursor="hand2",
                           anchor=tk.W)
        sub_lbl.pack(fill=tk.X)
        widgets.append(sub_lbl)

        status_lbl = tk.Label(text_col, text="未加载", font=("Microsoft YaHei UI", 7),
                              bg=self.CARD_BG, fg=self.MUTED, cursor="hand2",
                              anchor=tk.W)
        status_lbl.pack(fill=tk.X)
        widgets.append(status_lbl)

        for w in widgets:
            w.bind("<Enter>", on_enter)
            w.bind("<Leave>", on_leave)
            w.bind("<Button-1>", lambda _: on_click())
        return card, status_lbl

    def _make_btn(self, parent, text: str, command, bg=None, **kw):
        """统一样式的按钮"""
        b = tk.Button(parent, text=text, command=command,
                      bg=bg or self.PRIMARY, fg="white",
                      relief=tk.FLAT, padx=14, pady=6,
                      cursor="hand2", **kw)
        return b

    def _build_ui(self):
        # ---- 顶部标题 ----
        header = tk.Frame(self.root, bg="#1a1a2e", padx=20, pady=12)
        header.pack(fill=tk.X)
        tk.Label(header, text="Docx ↔ PDF 差异比对",
                 font=("Microsoft YaHei UI", 16, "bold"),
                 bg="#1a1a2e", fg="white").pack(side=tk.LEFT)
        tk.Label(header, text="OCR 驱动 · Git 风格展示",
                 font=("Microsoft YaHei UI", 9),
                 bg="#1a1a2e", fg="#a0a0c0").pack(side=tk.LEFT, padx=12)

        # ---- 拖放区域 2x2 网格 ----
        grid_frame = tk.Frame(self.root, bg=self.BG, padx=20, pady=10)
        grid_frame.pack(fill=tk.X)
        grid_frame.columnconfigure(0, weight=1)
        grid_frame.columnconfigure(1, weight=1)

        docx_card, self.lbl_docx = self._make_drop_zone(
            grid_frame, "拖入 DOCX", "原始 Word 文档（基准）", "📄",
            self.load_docx)
        docx_card.grid(row=0, column=0, sticky="nsew", padx=(0, 4), pady=(0, 4))

        pdf_card, self.lbl_pdf = self._make_drop_zone(
            grid_frame, "拖入 PDF", "扫描版 PDF（对比件）", "📑",
            self.load_pdf)
        pdf_card.grid(row=0, column=1, sticky="nsew", padx=(4, 0), pady=(0, 4))

        docx_folder_card, self.lbl_docx_folder = self._make_drop_zone(
            grid_frame, "选择 DOCX 文件夹", "批量 Word 文档目录", "📁",
            self._select_docx_folder)
        docx_folder_card.grid(row=1, column=0, sticky="nsew", padx=(0, 4), pady=(4, 0))

        pdf_folder_card, self.lbl_pdf_folder = self._make_drop_zone(
            grid_frame, "选择 PDF 文件夹", "批量 PDF 文件目录", "📂",
            self._select_pdf_folder)
        pdf_folder_card.grid(row=1, column=1, sticky="nsew", padx=(4, 0), pady=(4, 0))

        # ---- 操作栏 ----
        action_frame = tk.Frame(self.root, bg=self.BG, padx=20, pady=10)
        action_frame.pack(fill=tk.X)

        self._make_btn(action_frame, "分析差异 (Git 模式)",
                       self.analyze_diff).pack(side=tk.LEFT, padx=(0, 8))
        tk.Checkbutton(action_frame, text="比对符号", variable=self.compare_symbols,
                       command=lambda: self.analyze_diff(show_warning=False),
                       bg=self.BG, font=("Microsoft YaHei UI", 9)).pack(side=tk.LEFT, padx=8)
        tk.Checkbutton(action_frame, text="提示弹窗", variable=self.show_info,
                       bg=self.BG, font=("Microsoft YaHei UI", 9)).pack(side=tk.LEFT)
        self._make_btn(action_frame, "预览去红头 PDF", self.preview_cleaned_pdf,
                       bg="#e8a838").pack(side=tk.LEFT, padx=8)

        self._status_lbl = tk.Label(action_frame, text="",
                                    font=("Microsoft YaHei UI", 9),
                                    bg=self.BG, fg=self.MUTED)
        self._status_lbl.pack(side=tk.LEFT, padx=8)

        # 导出按钮（右侧）
        export_frame = tk.Frame(action_frame, bg=self.BG)
        export_frame.pack(side=tk.RIGHT)
        self._make_btn(export_frame, "导出：同步更新版 Docx", self.save_synced_docx,
                       bg=self.SUCCESS, font=("Microsoft YaHei UI", 9)).pack(
            side=tk.LEFT, padx=4)
        self._make_btn(export_frame, "导出：自动去红头版 Docx", self.save_dered_docx,
                       bg="#e8a838", font=("Microsoft YaHei UI", 9)).pack(
            side=tk.LEFT, padx=4)

        # ---- 批量导航栏（初始隐藏） ----
        self._batch_nav_frame = tk.Frame(self.root, bg=self.BG)
        # 不 pack，等批量模式激活时再显示
        self._make_btn(self._batch_nav_frame, "← 上一对", self._nav_prev,
                       bg="#6c757d", font=("Microsoft YaHei UI", 9)).pack(
            side=tk.LEFT, padx=(0, 4))
        self._batch_lbl = tk.Label(self._batch_nav_frame, text="第 1/1 对",
                                   font=("Microsoft YaHei UI", 9, "bold"),
                                   bg=self.BG, fg=self.TEXT)
        self._batch_lbl.pack(side=tk.LEFT, padx=8)
        self._make_btn(self._batch_nav_frame, "下一对 →", self._nav_next,
                       bg="#6c757d", font=("Microsoft YaHei UI", 9)).pack(
            side=tk.LEFT, padx=4)

        # ---- 差异展示区 ----
        mid_frame = tk.Frame(self.root, padx=20, pady=15, bg=self.BG)
        mid_frame.pack(fill=tk.BOTH, expand=True)

        tk.Label(mid_frame,
                 text="鼠标悬停更改行可操作  |  红底=删除  绿底=新增  |  浅色=已忽略  深色=已同意",
                 font=("Microsoft YaHei UI", 9),
                 bg=self.BG, fg=self.MUTED).pack(anchor=tk.W, pady=(0, 6))

        text_frame = tk.Frame(mid_frame, bg=self.CARD_BG,
                              highlightbackground=self.BORDER, highlightthickness=1)
        text_frame.pack(fill=tk.BOTH, expand=True)

        self.txt_diff = tk.Text(text_frame, font=("Consolas", 11), wrap=tk.WORD,
                                exportselection=True,
                                selectbackground=self.PRIMARY, selectforeground="white",
                                inactiveselectbackground="#7a9ec5",
                                bg=self.CARD_BG, fg=self.TEXT, bd=0,
                                padx=12, pady=10)
        self.txt_diff.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scroll = tk.Scrollbar(text_frame, command=self.txt_diff.yview, bg=self.BG)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.txt_diff.config(yscrollcommand=scroll.set)

        self.txt_diff.tag_config("del_on", background="#ffd4d4", foreground="#b30000",
                                 font=("Consolas", 11, "bold"))
        self.txt_diff.tag_config("add_on", background="#c8f0c8", foreground="#008000",
                                 font=("Consolas", 11, "bold"))
        self.txt_diff.tag_config("del_off", background="#fff5f5", foreground="#ccaaaa")
        self.txt_diff.tag_config("add_off", background="#f5fff5", foreground="#aaccaa")
        self.txt_diff.tag_config("header", foreground="#6a737d",
                                 font=("Consolas", 10, "italic"))
        self.txt_diff.tag_raise("sel")

        self._setup_copy_menu(self.txt_diff)

    def _setup_copy_menu(self, widget):
        """为Text组件添加右键复制菜单"""
        menu = tk.Menu(widget, tearoff=0)
        menu.add_command(label="复制", command=lambda: self._do_copy(widget))

        def on_right_click(e):
            try:
                self._saved_sel = widget.selection_get()
            except tk.TclError:
                self._saved_sel = None
            menu.post(e.x_root, e.y_root)

        widget.bind("<Button-3>", on_right_click)
        widget.bind("<Control-c>", lambda _: self._do_copy(widget))

    def _do_copy(self, widget):
        try:
            sel = getattr(self, '_saved_sel', None) or widget.selection_get()
            widget.clipboard_clear()
            widget.clipboard_append(sel)
        except tk.TclError:
            pass

    def _is_red_header(self, paragraph):
        """检查段落中的文字是否包含红色字体"""
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb:
                r, g, b = run.font.color.rgb
                if r > 200 and g < 60 and b < 60:
                    return True
        return False

    def _load_docx_path(self, path: str):
        """从路径加载 DOCX（支持拖放）"""
        # 非批量模式：清除批量状态
        if not self._pairs:
            self._batch_nav_frame.pack_forget()
            self._pair_data = {}
            self._pair_index = -1
        self.docx_path = path
        self.lbl_docx.config(text=os.path.basename(path), fg=self.PRIMARY)
        self.doc_obj = Document(path)
        body_paragraphs = []
        self._docx_paragraphs = []
        for p in self.doc_obj.paragraphs:
            if self._is_red_header(p):
                continue
            if p.text.strip():
                body_paragraphs.append(p.text)
                self._docx_paragraphs.append(p.text)
        raw_text = "\n".join(body_paragraphs)
        self.docx_text = self._normalize_text(raw_text)
        self.diff_blocks = []

    def load_docx(self):
        """通过文件对话框加载 DOCX"""
        path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if path:
            self._load_docx_path(path)

    def _scan_folder(self, path: str, exts: set):
        """扫描文件夹中指定扩展名的文件，返回 {stem: fullpath}"""
        result = {}
        for f in os.listdir(path):
            full = os.path.join(path, f)
            if not os.path.isfile(full):
                continue
            stem, ext = os.path.splitext(f)
            if ext.lower() in exts:
                result[stem] = full
        return result

    def _try_folder_match(self):
        """当两个文件夹都选定后，自动按文件名匹配并加载（支持批量）"""
        if not self._docx_folder or not self._pdf_folder:
            return

        docx_files = self._scan_folder(self._docx_folder, {".docx", ".doc"})
        pdf_files = self._scan_folder(self._pdf_folder, {".pdf"})

        pairs = []
        for stem in docx_files:
            if stem in pdf_files:
                pairs.append((docx_files[stem], pdf_files[stem]))

        if not pairs:
            self._alert("提示",
                f"两个文件夹中未找到文件名匹配的 DOCX/PDF 对。\n"
                f"DOCX 文件夹: {len(docx_files)} 个\n"
                f"PDF 文件夹: {len(pdf_files)} 个",
                "warn")
            return

        self._pairs = pairs
        self._pair_index = 0
        self._pair_data = {}

        # 显示批量导航
        if len(pairs) > 1:
            self._batch_nav_frame.pack(fill=tk.X, pady=(0, 5))

        docx_path, pdf_path = pairs[0]
        self._load_docx_path(docx_path)
        self._load_pdf_path(pdf_path, pair_index=0)
        if len(pairs) > 1:
            self._alert("提示",
                f"找到 {len(pairs)} 对匹配文件，请先裁剪第1对PDF。\n"
                f"完成后可选择是否复用裁剪区域。",
                "info")

    def _on_pair_ocr_done(self, index: int):
        """单对OCR完成：保存数据，若为首对则询问裁剪复用策略"""
        self._save_pair_data(index)
        self._update_batch_nav()
        if index == 0 and len(self._pairs) > 1:
            self._ask_crop_reuse()
        elif index > 0 and getattr(self, '_batch_crop_mode', '') == 'individual':
            self._batch_load_next_with_crop(index + 1)

    def _ask_crop_reuse(self):
        """弹窗询问是否复用首对裁剪区域"""
        win = tk.Toplevel(self.root)
        win.withdraw()
        win.title("批量裁剪策略")
        win.resizable(False, False)
        win.transient(self.root)
        win.grab_set()
        f = tk.Frame(win, padx=20, pady=15, bg="white")
        f.pack(fill=tk.BOTH, expand=True)
        tk.Label(f, text=f"第1对已处理完成，剩余 {len(self._pairs) - 1} 对。",
                 font=("Microsoft YaHei UI", 10), bg="white").pack(pady=(0, 5))
        tk.Label(f, text="是否将相同的裁剪区域应用于其余文件？",
                 font=("Microsoft YaHei UI", 10, "bold"), bg="white").pack(pady=(0, 12))
        btn_frame = tk.Frame(f, bg="white")
        btn_frame.pack()
        tk.Button(btn_frame, text="是，全部复用",
                  command=lambda: self._on_crop_choice(win, 'reuse'),
                  bg=self.PRIMARY, fg="white", relief=tk.FLAT,
                  padx=16, pady=6, cursor="hand2",
                  font=("Microsoft YaHei UI", 10)).pack(side=tk.LEFT, padx=4)
        tk.Button(btn_frame, text="否，逐对裁剪",
                  command=lambda: self._on_crop_choice(win, 'individual'),
                  bg="#6c757d", fg="white", relief=tk.FLAT,
                  padx=16, pady=6, cursor="hand2",
                  font=("Microsoft YaHei UI", 10)).pack(side=tk.LEFT, padx=4)
        win.bind("<Escape>", lambda _: self._on_crop_choice(win, 'reuse'))
        win.update_idletasks()
        pw, ph = self.root.winfo_width(), self.root.winfo_height()
        px, py = self.root.winfo_x(), self.root.winfo_y()
        ww, wh = win.winfo_width(), win.winfo_height()
        win.geometry(f"+{px + (pw - ww) // 2}+{py + (ph - wh) // 2}")
        win.deiconify()
        win.wait_window()

    def _on_crop_choice(self, win, mode: str):
        self._batch_crop_mode = mode
        win.destroy()
        if mode == 'reuse':
            self._batch_load_next(1)
        else:
            self._batch_load_next_with_crop(1)

    def _save_pair_data(self, index: int):
        self._pair_data[index] = {
            'docx_text': self.docx_text,
            'pdf_text': self.pdf_text,
            'doc_obj': self.doc_obj,
            'cleaned': self.cleaned_pdf_images,
            'docx_path': self.docx_path,
            'pdf_path': self.pdf_path,
            'pdf_name': self._pdf_name,
            'docx_paragraphs': self._docx_paragraphs,
        }

    def _update_batch_nav(self):
        total = len(self._pairs)
        loaded = len(self._pair_data)
        self._batch_lbl.config(text=f"第 {self._pair_index + 1}/{total} 对（已加载 {loaded} 对）")

    def _batch_load_next(self, index: int):
        """后台加载下一对（跳过裁剪对话框，复用首对 discard_boxes）"""
        if index >= len(self._pairs):
            return
        docx_path, pdf_path = self._pairs[index]
        import threading

        def worker():
            try:
                # 加载 DOCX
                doc = Document(docx_path)
                paras = []
                for p in doc.paragraphs:
                    if not self._is_red_header(p) and p.text.strip():
                        paras.append(p.text)
                docx_text = self._normalize_text("\n".join(paras))

                # 加载 PDF 页面
                pdf_doc = fitz.open(pdf_path)
                mat = fitz.Matrix(1.5, 1.5)
                images = []
                for i in range(len(pdf_doc)):
                    page = pdf_doc.load_page(i)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data)).convert("RGB")
                    cleaned = self._remove_red_pixels(img)
                    # 应用首对的裁剪区域
                    boxes = self.discard_boxes[i % 2]
                    if boxes:
                        draw = ImageDraw.Draw(cleaned)
                        for box in boxes:
                            draw.rectangle(box, fill="white")
                    images.append(cleaned)

                # OCR
                pdf_lines = []
                for img in images:
                    lines = self._ocr_single(img)
                    pdf_lines.extend(lines)
                pdf_text = self._normalize_text("\n".join(pdf_lines))

                self._pair_data[index] = {
                    'docx_text': docx_text,
                    'pdf_text': pdf_text,
                    'doc_obj': doc,
                    'cleaned': images,
                    'docx_path': docx_path,
                    'pdf_path': pdf_path,
                    'pdf_name': os.path.basename(pdf_path),
                    'docx_paragraphs': paras,
                }
            except Exception as e:
                self._pair_data[index] = {'error': str(e)}

        def poll():
            if thread.is_alive():
                self.root.after(200, poll)
            else:
                if 'error' in self._pair_data.get(index, {}):
                    err = self._pair_data[index]['error']
                    self._alert("批量加载错误",
                                f"第{index + 1}对加载失败:\n{err}", "error")
                self._update_batch_nav()
                # 继续加载下一对
                self._batch_load_next(index + 1)

        thread = threading.Thread(target=worker, daemon=True)
        thread.start()
        self.root.after(100, poll)

    def _batch_load_next_with_crop(self, index: int):
        """逐对加载（含裁剪对话框）"""
        if index >= len(self._pairs):
            self._update_batch_nav()
            return
        docx_path, pdf_path = self._pairs[index]
        # 保存当前对的数据（如果还在激活对的话）
        if self._pair_index >= 0 and self._pair_index not in self._pair_data:
            self._save_pair_data(self._pair_index)
        self._load_docx_path(docx_path)
        self._load_pdf_path(pdf_path, pair_index=index)

    def _switch_to_pair(self, index: int):
        """切换到指定索引的比对对"""
        if index not in self._pair_data:
            return
        data = self._pair_data[index]
        self._pair_index = index
        self.docx_text = data['docx_text']
        self.pdf_text = data['pdf_text']
        self.doc_obj = data['doc_obj']
        self.cleaned_pdf_images = data['cleaned']
        self.docx_path = data['docx_path']
        self.pdf_path = data['pdf_path']
        self._pdf_name = data['pdf_name']
        self._docx_paragraphs = data.get('docx_paragraphs', [])
        self.diff_blocks = []

        self.lbl_docx.config(text=os.path.basename(self.docx_path), fg=self.PRIMARY)
        self.lbl_pdf.config(text=f"{self._pdf_name}（OCR完成）", fg="green")
        self._update_batch_nav()

        # 如果差异已展示，自动刷新
        if self.txt_diff.get("1.0", tk.END).strip():
            self.analyze_diff(show_warning=False)

    def _nav_prev(self):
        if self._pair_index > 0:
            self._switch_to_pair(self._pair_index - 1)

    def _nav_next(self):
        if self._pair_index < len(self._pairs) - 1:
            self._switch_to_pair(self._pair_index + 1)

    def _select_docx_folder(self):
        path = filedialog.askdirectory(title="选择存放 DOCX 文件的文件夹")
        if path:
            self._docx_folder = path
            self.lbl_docx_folder.config(text=os.path.basename(path), fg=self.PRIMARY)
            self._try_folder_match()

    def _select_pdf_folder(self):
        path = filedialog.askdirectory(title="选择存放 PDF 文件的文件夹")
        if path:
            self._pdf_folder = path
            self.lbl_pdf_folder.config(text=os.path.basename(path), fg=self.PRIMARY)
            self._try_folder_match()

    @staticmethod
    def _remove_red_pixels(img):
        """将偏红色像素替换为白色（numpy 向量化，比逐像素循环快数十倍）"""
        arr = np.array(img)
        r, g, b = arr[:, :, 0], arr[:, :, 1], arr[:, :, 2]
        mask = r > np.maximum(g, b) + 30
        arr[mask] = [255, 255, 255]
        return Image.fromarray(arr)

    @staticmethod
    def _normalize_text(text):
        """归一化文本：全角转半角、合并连续空格、去除行首行尾空白"""
        result = []
        for ch in text:
            code = ord(ch)
            if 0xFF01 <= code <= 0xFF5E:  # 全角ASCII → 半角
                result.append(chr(code - 0xFEE0))
            elif code == 0x3000:  # 全角空格 → 半角空格
                result.append(' ')
            else:
                result.append(ch)
        normalized = ''.join(result)
        # 合并连续空格，逐行strip
        lines = []
        for line in normalized.splitlines():
            cleaned = ' '.join(line.split())
            if cleaned:
                lines.append(cleaned)
        return '\n'.join(lines)

    def _load_pdf_path(self, path: str, pair_index: int = None):
        """从路径加载 PDF（支持拖放），包括去红头 + 裁剪 → OCR"""
        self.pdf_path = path
        self._pdf_name = os.path.basename(path)
        self._set_status(f"正在加载 {self._pdf_name} ...", self.WARN)

        self.cleaned_pdf_images = []
        self.crop_box = None
        self._current_pair_index = pair_index
        try:
            doc = fitz.open(path)
            mat = fitz.Matrix(1.5, 1.5)
            for i in range(len(doc)):
                page = doc.load_page(i)
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data)).convert("RGB")
                cleaned = self._remove_red_pixels(img)
                self.cleaned_pdf_images.append(cleaned)
                self._set_status(f"加载中... 第 {i + 1}/{len(doc)} 页")

            self.lbl_pdf.config(text=f"{self._pdf_name}（已加载，请裁剪）", fg=self.WARN)
            self._show_crop_dialog()
        except Exception as e:
            self.lbl_pdf.config(text=f"{self._pdf_name}（加载失败）", fg="red")
            self._set_status("")
            self._alert("错误", f"加载PDF异常:\n{str(e)}", "error")

    def load_pdf(self):
        """通过文件对话框加载 PDF"""
        path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
        if path:
            self._load_pdf_path(path)

    def _show_crop_dialog(self):
        """两步裁剪：第1页（单数页基准）→ 第2页（双数页基准），每页可框选多个区域"""
        if not self.cleaned_pdf_images:
            return

        self.discard_boxes: list = [[], []]  # 每页一个列表，存放多个 (x1,y1,x2,y2)
        win = tk.Toplevel(self.root)
        win.title("裁剪PDF（框选要丢弃的区域，可多选）")
        win.geometry("900x700")

        info_lbl = tk.Label(win, text="", fg="blue")
        info_lbl.pack(pady=5)

        frame = tk.Frame(win)
        frame.pack(fill=tk.BOTH, expand=True)
        v_scroll = tk.Scrollbar(frame, orient=tk.VERTICAL)
        v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        h_scroll = tk.Scrollbar(frame, orient=tk.HORIZONTAL)
        h_scroll.pack(side=tk.BOTTOM, fill=tk.X)
        canvas = tk.Canvas(frame, cursor="cross", bg="gray",
                           yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_scroll.config(command=canvas.yview)
        h_scroll.config(command=canvas.xview)

        def load_page(page_idx):
            """加载指定页到画布，返回 (scale, offset_x, offset_y)"""
            canvas.update()
            cw = max(canvas.winfo_width(), 100)
            ch = max(canvas.winfo_height(), 100)
            img = self.cleaned_pdf_images[page_idx]
            s = min(cw / img.width, ch / img.height, 1.0)
            dw = int(img.width * s)
            dh = int(img.height * s)
            display = img.resize((dw, dh), Image.Resampling.LANCZOS)
            tk_img = ImageTk.PhotoImage(display)
            canvas.__dict__['_img'] = tk_img
            canvas.delete("all")
            ox = cw // 2 - dw // 2
            oy = ch // 2 - dh // 2
            canvas.create_image(cw // 2, ch // 2, anchor=tk.CENTER, image=tk_img)
            canvas.config(scrollregion=canvas.bbox(tk.ALL))
            # 重绘已保存的框选区域
            for box in self.discard_boxes[page_idx]:
                x1, y1, x2, y2 = box
                canvas.create_rectangle(
                    ox + x1 * s, oy + y1 * s, ox + x2 * s, oy + y2 * s,
                    outline="blue", width=2, dash=(4, 4),
                )
            return s, ox, oy

        step = 0  # 0 = 设置第1页, 1 = 设置第2页
        rect_id = None
        start_x, start_y = 0, 0
        cur_scale = 1.0
        img_offset_x, img_offset_y = 0, 0

        def refresh_canvas():
            nonlocal rect_id, start_x, start_y, cur_scale, img_offset_x, img_offset_y
            rect_id = None
            start_x, start_y = 0, 0
            cur_scale, img_offset_x, img_offset_y = load_page(step)
            count = len(self.discard_boxes[step])
            suffix = f"（已选 {count} 个区域）" if count else ""
            if step == 0:
                info_lbl.config(text=f"第1步：框选第1页中要丢弃的区域{suffix}（单数页使用此位置）")
                btn_next.config(text="保存并设置第2页 →", state=tk.NORMAL)
                btn_skip.config(text="跳过此步（不擦除单数页）", state=tk.NORMAL)
                btn_undo.config(state=tk.NORMAL if count else tk.DISABLED)
            else:
                info_lbl.config(text=f"第2步：框选第2页中要丢弃的区域{suffix}（双数页使用此位置）")
                btn_next.config(text="确认并开始OCR", state=tk.NORMAL)
                btn_skip.config(text="跳过（不擦除双数页）", state=tk.NORMAL)
                btn_undo.config(state=tk.NORMAL if count else tk.DISABLED)

        # 鼠标滚轮
        def on_wheel(event):
            canvas.yview_scroll(-1 if event.delta > 0 else 1, tk.UNITS)
        canvas.bind("<MouseWheel>", on_wheel)

        # 框选交互
        def on_down(event):
            nonlocal start_x, start_y, rect_id
            start_x = canvas.canvasx(event.x)
            start_y = canvas.canvasy(event.y)
            if rect_id is not None:
                canvas.delete(rect_id)
            rect_id = canvas.create_rectangle(
                start_x, start_y, start_x, start_y,
                outline="red", width=2, dash=(4, 4),
            )

        def on_drag(event):
            if rect_id is not None:
                cur_x = canvas.canvasx(event.x)
                cur_y = canvas.canvasy(event.y)
                canvas.coords(rect_id, start_x, start_y, cur_x, cur_y)

        def on_up(event):
            nonlocal rect_id
            end_x = canvas.canvasx(event.x)
            end_y = canvas.canvasy(event.y)
            x1, x2 = sorted([start_x, end_x])
            y1, y2 = sorted([start_y, end_y])
            if (x2 - x1) > 5 and (y2 - y1) > 5:
                raw = (
                    int((x1 - img_offset_x) / cur_scale),
                    int((y1 - img_offset_y) / cur_scale),
                    int((x2 - img_offset_x) / cur_scale),
                    int((y2 - img_offset_y) / cur_scale),
                )
                self.discard_boxes[step].append(raw)
                # 把当前矩形从临时色改为持久色
                if rect_id is not None:
                    canvas.itemconfig(rect_id, outline="blue")
                rect_id = None
                refresh_canvas()
            else:
                if rect_id is not None:
                    canvas.delete(rect_id)
                    rect_id = None

        canvas.bind("<ButtonPress-1>", on_down)
        canvas.bind("<B1-Motion>", on_drag)
        canvas.bind("<ButtonRelease-1>", on_up)

        btn_frame = tk.Frame(win)
        btn_frame.pack(pady=8)

        def on_undo():
            if self.discard_boxes[step]:
                self.discard_boxes[step].pop()
                refresh_canvas()

        def on_next():
            nonlocal step
            if step == 0:
                step = 1
                refresh_canvas()
            else:
                win.destroy()
                pi = getattr(self, '_current_pair_index', None)
                self._apply_crop_and_ocr(pair_index=pi)

        def on_skip():
            nonlocal step
            self.discard_boxes[step] = []
            if step == 0:
                step = 1
                refresh_canvas()
            else:
                win.destroy()
                pi = getattr(self, '_current_pair_index', None)
                self._apply_crop_and_ocr(pair_index=pi)

        btn_undo = tk.Button(btn_frame, text="撤销上一个区域", command=on_undo, width=18,
                             state=tk.DISABLED)
        btn_undo.pack(side=tk.LEFT, padx=10)
        btn_skip = tk.Button(btn_frame, text="", command=on_skip, width=20)
        btn_skip.pack(side=tk.LEFT, padx=10)
        btn_next = tk.Button(btn_frame, text="", command=on_next, bg="#d1ecf1", width=22)
        btn_next.pack(side=tk.LEFT, padx=10)

        refresh_canvas()

    @staticmethod
    def _ocr_single(img):
        """对单张图片执行OCR"""
        result, _ = ocr_engine(np.array(img))
        lines = []
        if result:
            for _, text, _ in result:
                if text.strip():
                    lines.append(text.strip())
        return lines

    def _apply_crop_and_ocr(self, pair_index: int = None):
        """异步执行裁剪+OCR，不阻塞UI。pair_index 非 None 时表示批量模式"""
        self._set_status("正在准备OCR...", "orange")

        # 先同步擦除框选区域（轻量操作）
        for i, img in enumerate(self.cleaned_pdf_images):
            boxes = self.discard_boxes[i % 2]
            if boxes:
                draw = ImageDraw.Draw(img)
                for box in boxes:
                    draw.rectangle(box, fill="white")

        # OCR 放到后台线程，UI 线程用 after 轮询进度
        import threading
        total = len(self.cleaned_pdf_images)
        results: list = []
        progress = {"current": 0}
        error: list = [None]

        def worker():
            try:
                for i, img in enumerate(self.cleaned_pdf_images):
                    lines = self._ocr_single(img)
                    results.append((i, lines))
                    progress["current"] = i + 1
            except Exception as e:
                error[0] = e

        thread = threading.Thread(target=worker, daemon=True)
        thread.start()

        def poll():
            cur = progress["current"]
            suffix = f" (第{pair_index + 1}/{len(self._pairs)}对)" if pair_index is not None and self._pairs else ""
            self._set_status(f"OCR中... 第 {cur}/{total} 页{suffix}", "orange")
            if thread.is_alive():
                self.root.after(150, poll)
            elif error[0]:
                self.lbl_pdf.config(text=f"{self._pdf_name}（OCR失败）", fg="red")
                self._alert("OCR错误", f"解析PDF异常:\n{str(error[0])}", "error")
                self._set_status("")
            else:
                results.sort(key=lambda x: x[0])
                pdf_lines = []
                for _, lines in results:
                    pdf_lines.extend(lines)
                raw_pdf_text = "\n".join(pdf_lines)
                self.pdf_text = self._normalize_text(raw_pdf_text)
                self.lbl_pdf.config(text=f"{self._pdf_name}（OCR完成）", fg="green")
                self._set_status("")
                if pair_index is not None:
                    self._on_pair_ocr_done(pair_index)
                else:
                    self._alert("完成", "PDF处理完成，现在可以点击「分析差异」进行比对。", "info")

        self.root.after(100, poll)

    def analyze_diff(self, show_warning=True):
        if not self.docx_text or not self.pdf_text:
            if show_warning:
                self._alert("提示", "请同时加载Docx和PDF文件后再进行比对。", "warn")
            return

        if self.compare_symbols.get():
            docx_flat = ''.join(ch for ch in self.docx_text if not ch.isspace())
            pdf_flat = ''.join(ch for ch in self.pdf_text if not ch.isspace())
        else:
            docx_flat = ''.join(
                ch for ch in self.docx_text
                if not ch.isspace() and (
                    '一' <= ch <= '鿿' or '㐀' <= ch <= '䶿'
                    or 'a' <= ch <= 'z' or 'A' <= ch <= 'Z'
                    or '0' <= ch <= '9'
                ))
            pdf_flat = ''.join(
                ch for ch in self.pdf_text
                if not ch.isspace() and (
                    '一' <= ch <= '鿿' or '㐀' <= ch <= '䶿'
                    or 'a' <= ch <= 'z' or 'A' <= ch <= 'Z'
                    or '0' <= ch <= '9'
                ))

        # 构建位置映射表（flat → full_text 位置），用于后续合并
        full_text = '\n'.join(self._docx_paragraphs)
        self._docx_flat_positions = []
        for i, ch in enumerate(full_text):
            if self.compare_symbols.get():
                if not ch.isspace():
                    self._docx_flat_positions.append(i)
            else:
                if not ch.isspace() and (
                    '一' <= ch <= '鿿' or '㐀' <= ch <= '䶿'
                    or 'a' <= ch <= 'z' or 'A' <= ch <= 'Z'
                    or '0' <= ch <= '9'
                ):
                    self._docx_flat_positions.append(i)

        matcher = difflib.SequenceMatcher(None, docx_flat, pdf_flat)
        self.diff_blocks = []
        block_id = 0
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            old_text = docx_flat[i1:i2]
            new_text = pdf_flat[j1:j2]
            if tag == 'equal':
                self.diff_blocks.append({
                    'tag': tag, 'old': old_text, 'new': '',
                    'accepted': True, 'id': block_id,
                })
            elif tag in ('delete', 'insert', 'replace'):
                self.diff_blocks.append({
                    'tag': tag, 'old': old_text, 'new': new_text,
                    'accepted': False, 'id': block_id,
                })
            block_id += 1

        self._hide_hover_popup()
        self._render_diff()
        self.txt_diff.bind("<Motion>", self._on_diff_motion)

        if show_warning:
            self._alert("分析完成",
                "鼠标悬停在更改行上可弹出操作窗口。\n导出时将仅并入已同意的更改。",
                "info")

    def _render_diff(self):
        """根据 self.diff_blocks 渲染差异文本"""
        # 保存第一可见行的行号
        try:
            top_line = int(str(self.txt_diff.index("@0,0")).split(".")[0])
        except (ValueError, IndexError):
            top_line = 1
        self.txt_diff.delete("1.0", tk.END)
        mode = "（比对符号）" if self.compare_symbols.get() else "（忽略符号）"
        pair_info = ""
        if self._pairs and self._pair_index >= 0:
            pair_info = f"  |  第 {self._pair_index + 1}/{len(self._pairs)} 对"
        approved = sum(1 for b in self.diff_blocks if b['accepted'] and b['tag'] != 'equal')
        total = sum(1 for b in self.diff_blocks if b['tag'] != 'equal')
        self.txt_diff.insert(tk.END,
            f"--- DOCX（基准）{mode}{pair_info}  |  已同意 {approved}/{total} 处更改\n", "header")
        self.txt_diff.insert(tk.END, "+++ PDF（对比件）\n", "header")

        for b in self.diff_blocks:
            tag = b['tag']
            bid = b['id']

            if tag == 'equal':
                self.txt_diff.insert(tk.END, f"  {b['old']}\n", f"block_{bid}")
            elif tag == 'delete':
                t = "del_on" if b['accepted'] else "del_off"
                self.txt_diff.insert(tk.END, f"- {b['old']}\n",
                                     (f"block_{bid}", t))
            elif tag == 'insert':
                t = "add_on" if b['accepted'] else "add_off"
                self.txt_diff.insert(tk.END, f"+ {b['new']}\n",
                                     (f"block_{bid}", t))
            elif tag == 'replace':
                dt = "del_on" if b['accepted'] else "del_off"
                at = "add_on" if b['accepted'] else "add_off"
                self.txt_diff.insert(tk.END, f"- {b['old']}\n",
                                     (f"block_{bid}", dt))
                self.txt_diff.insert(tk.END, f"+ {b['new']}\n",
                                     (f"block_{bid}", at))

        # 恢复滚动位置
        self.txt_diff.update_idletasks()
        self.txt_diff.see(f"{top_line}.0")

    def _get_block_at(self, x, y):
        """返回鼠标位置所在的 diff 块（仅非 equal 块），无则返回 None"""
        idx = self.txt_diff.index(f"@{x},{y}")
        tags = self.txt_diff.tag_names(idx)
        for tag in tags:
            if tag.startswith("block_"):
                try:
                    bid = int(tag.split("_")[1])
                    for b in self.diff_blocks:
                        if b['id'] == bid and b['tag'] != 'equal':
                            return b
                except (ValueError, IndexError):
                    pass
        return None

    def _on_diff_motion(self, event):
        """鼠标移动：检测悬停的更改块，弹出操作悬浮窗"""
        block = self._get_block_at(event.x, event.y)
        if block is None:
            self._hide_hover_popup()
            return
        if block['id'] == self._hover_block_id:
            return  # 同一块，不刷新

        self._hover_block_id = block['id']
        self._show_hover_popup(block, event)

    def _show_hover_popup(self, block, event):
        """在鼠标附近显示操作悬浮窗"""
        self._hide_hover_popup()

        popup = tk.Toplevel(self.root)
        popup.overrideredirect(True)
        popup.attributes("-topmost", True)
        popup.configure(bg="#2d2d2d")

        # 内容
        inner = tk.Frame(popup, bg="#2d2d2d", padx=10, pady=8)
        inner.pack()

        tag_label = {"delete": "删除", "insert": "新增", "replace": "替换"}
        label = tag_label.get(block['tag'], '变更')
        status = "已同意" if block['accepted'] else "已忽略"

        tk.Label(inner, text=f"{label} | {status}",
                 font=("Microsoft YaHei UI", 9, "bold"),
                 bg="#2d2d2d", fg="white").pack(pady=(0, 4))

        # 变更内容预览（截取前30字符）
        if block['tag'] in ('delete', 'replace'):
            preview = block['old'][:30] + ('...' if len(block['old']) > 30 else '')
            tk.Label(inner, text=f"- {preview}",
                     font=("Consolas", 9), bg="#2d2d2d",
                     fg="#ff8888", wraplength=300, justify=tk.LEFT).pack(anchor=tk.W)
        if block['tag'] in ('insert', 'replace'):
            preview = block['new'][:30] + ('...' if len(block['new']) > 30 else '')
            tk.Label(inner, text=f"+ {preview}",
                     font=("Consolas", 9), bg="#2d2d2d",
                     fg="#88ff88", wraplength=300, justify=tk.LEFT).pack(anchor=tk.W)

        # 切换按钮
        btn_text = "改为忽略" if block['accepted'] else "改为同意"
        btn_bg = self.PRIMARY if not block['accepted'] else "#6c757d"
        tk.Button(inner, text=btn_text,
                  command=lambda b=block: self._do_toggle(b),
                  bg=btn_bg, fg="white", relief=tk.FLAT,
                  padx=12, pady=3, cursor="hand2",
                  font=("Microsoft YaHei UI", 9)).pack(pady=(6, 0))

        # 定位：鼠标右下方偏移
        x = self.root.winfo_x() + event.widget.winfo_rootx() - self.root.winfo_rootx() + event.x + 15
        y = self.root.winfo_y() + event.widget.winfo_rooty() - self.root.winfo_rooty() + event.y + 10
        popup.geometry(f"+{x}+{y}")
        self._hover_popup = popup

    def _hide_hover_popup(self):
        if self._hover_popup:
            try:
                self._hover_popup.destroy()
            except tk.TclError:
                pass
            self._hover_popup = None
            self._hover_block_id = -1

    def _do_toggle(self, block):
        """切换块状态并刷新"""
        block['accepted'] = not block['accepted']
        self._hover_block_id = -1  # 强制刷新悬浮窗
        self._render_diff()

    def preview_cleaned_pdf(self):
        """在新窗口中预览去红头后的PDF图像"""
        if not self.cleaned_pdf_images:
            self._alert("提示", "请先加载PDF文件。", "warn")
            return

        win = tk.Toplevel(self.root)
        win.title("去红头PDF预览")
        img_w, img_h = self.cleaned_pdf_images[0].size
        screen_w = self.root.winfo_screenwidth() - 100
        screen_h = self.root.winfo_screenheight() - 150
        scale = min(screen_w / img_w, screen_h / img_h, 1.0)
        win_w = max(int(img_w * scale), 300)
        win_h = max(int(img_h * scale) + 50, 200)
        win.geometry(f"{win_w}x{win_h}")

        # 翻页控制
        nav_frame = tk.Frame(win, pady=5)
        nav_frame.pack(side=tk.TOP, fill=tk.X)

        # 画布
        frame = tk.Frame(win)
        frame.pack(fill=tk.BOTH, expand=True)
        v_scroll = tk.Scrollbar(frame, orient=tk.VERTICAL)
        v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        h_scroll = tk.Scrollbar(frame, orient=tk.HORIZONTAL)
        h_scroll.pack(side=tk.BOTTOM, fill=tk.X)
        canvas = tk.Canvas(frame, cursor="cross", bg="gray",
                           yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        v_scroll.config(command=canvas.yview)
        h_scroll.config(command=canvas.xview)

        lbl_page = tk.Label(nav_frame, text="")
        lbl_page.pack(side=tk.LEFT, padx=10)
        img_idx = 0

        def show_page():
            # 动态适配画布大小居中显示
            canvas.update()
            cw = max(canvas.winfo_width(), 100)
            ch = max(canvas.winfo_height(), 100)
            img = self.cleaned_pdf_images[img_idx]
            scale = min(cw / img.width, ch / img.height, 1.0)
            w = int(img.width * scale)
            h = int(img.height * scale)
            tk_img = ImageTk.PhotoImage(img.resize((w, h), Image.Resampling.LANCZOS))
            canvas.__dict__['_img'] = tk_img
            canvas.delete("all")
            canvas.create_image(cw // 2, ch // 2, anchor=tk.CENTER, image=tk_img)
            canvas.config(scrollregion=canvas.bbox(tk.ALL))
            lbl_page.config(text=f"第 {img_idx + 1} / {len(self.cleaned_pdf_images)} 页 (缩放 {scale:.0%})")

        def on_wheel(event):
            canvas.yview_scroll(-1 if event.delta > 0 else 1, tk.UNITS)
        canvas.bind("<MouseWheel>", on_wheel)

        def prev():
            nonlocal img_idx
            if img_idx > 0:
                img_idx -= 1
                show_page()

        def next_page():
            nonlocal img_idx
            if img_idx < len(self.cleaned_pdf_images) - 1:
                img_idx += 1
                show_page()

        tk.Button(nav_frame, text="上一页", command=prev).pack(side=tk.LEFT, padx=5)
        tk.Button(nav_frame, text="下一页", command=next_page).pack(side=tk.LEFT, padx=5)
        show_page()

    def save_synced_docx(self):
        """导出：将已同意的更改并入 DOCX（未同意则保留原文）"""
        if not self.doc_obj:
            return
        save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                   filetypes=[("Word Documents", "*.docx")])
        if not save_path:
            return

        # 检查是否有已同意的更改
        accepted = [b for b in self.diff_blocks if b['accepted'] and b['tag'] != 'equal']
        if not accepted:
            self.doc_obj.save(save_path)
            self._alert("成功", "未同意任何更改，已导出原版文档。", "info")
            return

        # 重建 accepted 目标 flat text
        target_parts = []
        for b in self.diff_blocks:
            if b['tag'] == 'equal':
                target_parts.append(b['old'])
            elif b['tag'] == 'delete':
                if not b['accepted']:
                    target_parts.append(b['old'])
            elif b['tag'] == 'insert':
                if b['accepted']:
                    target_parts.append(b['new'])
            elif b['tag'] == 'replace':
                target_parts.append(b['new'] if b['accepted'] else b['old'])
        target_flat = ''.join(target_parts)

        # 重建原始 flat text 及位置映射
        full_text = '\n'.join(self._docx_paragraphs)
        original_flat = []
        flat_to_full = []  # flat index → position in full_text
        for i, ch in enumerate(full_text):
            if self.compare_symbols.get():
                if not ch.isspace():
                    original_flat.append(ch)
                    flat_to_full.append(i)
            else:
                if not ch.isspace() and (
                    '一' <= ch <= '鿿' or '㐀' <= ch <= '䶿'
                    or 'a' <= ch <= 'z' or 'A' <= ch <= 'Z'
                    or '0' <= ch <= '9'
                ):
                    original_flat.append(ch)
                    flat_to_full.append(i)
        original_flat = ''.join(original_flat)

        # Diff original → target，得到 flat 空间的操作码
        matcher = difflib.SequenceMatcher(None, original_flat, target_flat)
        edits = []  # [(full_start, delete_len, insert_str), ...]
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                continue
            fs = flat_to_full[i1] if i1 < len(flat_to_full) else len(full_text)
            fe = flat_to_full[i2] if i2 < len(flat_to_full) else len(full_text)
            dlen = fe - fs
            insert = target_flat[j1:j2] if tag in ('insert', 'replace') else ''
            edits.append((fs, dlen, insert))

        # 逆序应用编辑到 full_text（避免位置偏移）
        edits.sort(key=lambda x: x[0], reverse=True)
        chars = list(full_text)
        for fs, dlen, insert in edits:
            del chars[fs:fs + dlen]
            for ch in reversed(insert):
                chars.insert(fs, ch)
        modified = ''.join(chars)

        # 按段落边界回写到 DOCX
        new_paras = modified.split('\n')
        body_paras = [p for p in self.doc_obj.paragraphs
                      if not self._is_red_header(p) and p.text.strip()]
        for i, para_text in enumerate(new_paras):
            if i < len(body_paras):
                p = body_paras[i]
                if p.runs:
                    p.runs[0].text = para_text
                    for r in p.runs[1:]:
                        r.text = ''
                else:
                    p.text = para_text
            else:
                # 新增段落
                new_p = body_paras[-1].insert_paragraph_after(para_text)
                if new_p.runs:
                    new_p.runs[0].text = para_text
        # 多余的原段落清空
        for p in body_paras[len(new_paras):]:
            p.text = ''

        self.doc_obj.save(save_path)
        self._alert("成功",
            f"已并入 {len(accepted)} 处更改并导出文档。", "info")

    def save_dered_docx(self):
        """自动去红头：扫描前部段落，检测红色字体并清除"""
        if not self.doc_obj:
            return

        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if not save_path:
            return

        # 创建一个用于修改的副本对象
        export_doc = Document(self.docx_path)

        # 自动化去红头策略：
        # 1. 扫描前5个段落。
        # 2. 如果某个段落中存在文字颜色为红色（RGB: >200, <50, <50），
        #    或者包含政府公文红头常见字（如"文件"），则判定为红头，将其内容清空。
        paragraphs_to_check = export_doc.paragraphs[:6]

        for p in paragraphs_to_check:
            is_red_head = False
            # 检查段落中的每个文字运行块（Run）的颜色
            for run in p.runs:
                if run.font.color and run.font.color.rgb:
                    r, g, b = run.font.color.rgb
                    if r > 200 and g < 60 and b < 60:  # 判定为红色系字体
                        is_red_head = True
                        break

            # 如果命中了红头特征，清除该段落文本（保留位置占位，不破坏后续排版结构）
            if is_red_head or "文件" in p.text:
                p.text = ""
                # 连带清除可能遗留的下划线/边框属性
                p.paragraph_format.space_before = 0
                p.paragraph_format.space_after = 0

        try:
            export_doc.save(save_path)
            self._alert("成功", "自动去红头版Docx文档已成功导出。", "info")
        except Exception as e:
            self._alert("错误", f"保存失败: {str(e)}", "error")


if __name__ == "__main__":
    root = tk.Tk()
    app = DocxPdfReviewer(root)
    root.mainloop()
